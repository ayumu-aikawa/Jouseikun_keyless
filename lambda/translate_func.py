import io
import os
import boto3
import json
import docx
from datetime import datetime

# Imports the Google Cloud client library
# from google.cloud import vision
output_bucket = ''

    
def lambda_handler(event, context):

    # Instantiates a client
    # client = vision.ImageAnnotatorClient()
    translate = boto3.client('translate')
    


    # s3から画像取得
    s3 = boto3.client('s3')
    input_bucket=event['Records'][0]['s3']['bucket']['name']
    input_data=event['Records'][0]['s3']['object']['key']
    #拡張子チェック
    file_name = input_data.rsplit('.', 1)
        
    input_list=input_data.split('_')
    print(input_list)
    
    input_key=input_list[2]
    #拡張子ごとのファイル処理の分岐
    if file_name[1] == 'txt':
        content = s3.get_object(Bucket=input_bucket, Key=input_data)['Body'].read()
        input_text = content.decode('utf-8')
    elif file_name[1] == 'docx':
        content = ''
        
        obj_data = s3.get_object(Bucket=input_bucket, Key=input_data)['Body']
        doc = docx.Document(io.BytesIO(obj_data.read()))
        content = ''
        for para in doc.paragraphs:
            content += (para.text + '\n') 
        input_text = content
        print('docx 1 ok')
    
    if input_list[0]!='none':
    
        response = translate.translate_text(
            Text=input_text,
            SourceLanguageCode=input_list[0],
            TargetLanguageCode=input_list[1]
        )
        
        #翻訳結果をoutput_textに入れる
        output_text =response.get('TranslatedText')

        #拡張子ごとの翻訳結果の処理の分岐
        if file_name[1] == 'txt':
            #txtはそのまんま
            file_contents = output_text
        elif file_name[1] == 'docx':
            new_doc = docx.Document()
            new_doc.add_paragraph(output_text)
            new_doc.save('/tmp/' + input_key)
            file_contents = open('/tmp/' + input_key,'rb+')
            print('docx 2 ok')
    
    else:
        output_text = input_text
    
    print('ok1')
    # S3に書き込み
    s3_resource = boto3.resource('s3')
    output_key = f"{input_key}"
    print('ok2')
    
    obj = s3_resource.Object(output_bucket,output_key)
    obj.put( Body=file_contents )
    print('ok3')
    
    s3 = boto3.client('s3')
    s3.delete_object(Bucket=input_bucket, Key=input_data)
    print('ok4')

    return {
        'statusCode': 200
    }